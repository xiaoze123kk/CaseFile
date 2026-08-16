"""Unit tests for document text extraction and block splitting."""

from io import BytesIO

import pytest
from casefile.importers.text_extraction import (
    MAX_FILE_BYTES,
    ExtractionError,
    extract_text,
    split_blocks,
)
from docx import Document


def test_extract_text_plain_and_markdown():
    assert extract_text("a.txt", b"\xe6\x9c\xac\xe4\xbd\x93") == "本体"
    markdown = b"# \xe6\xa0\x87\xe9\xa2\x98\n\n- \xe9\xa1\xb9\xe7\x9b\xae\n"
    assert "标题" in extract_text("a.md", markdown)


def test_extract_text_docx():
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    document.save(buffer)
    text = extract_text("a.docx", buffer.getvalue())
    assert "第一段" in text and "第二段" in text


def test_extract_text_pdf():
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buffer = BytesIO()
    writer.write(buffer)
    text = extract_text("a.pdf", buffer.getvalue())
    assert text == ""  # blank page has no text


def test_extract_text_rejects_oversize():
    with pytest.raises(ExtractionError, match="5MB"):
        extract_text("a.txt", b"x" * (MAX_FILE_BYTES + 1))


def test_extract_text_rejects_unknown_extension():
    with pytest.raises(ExtractionError, match="不支持"):
        extract_text("a.exe", b"data")


def test_split_blocks():
    assert split_blocks("第一段\n\n\n第二段\n\n") == ["第一段", "第二段"]
    assert split_blocks("") == []
    assert split_blocks("只有一段") == ["只有一段"]
